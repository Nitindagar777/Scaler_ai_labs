"""Leak check v2: verify no original PII remains in the redacted document."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"Red Herring Prospectus_REDACTED.docx")

# Check for any remaining original PII
originals = [
    'Kushal Subbayya', 'Kushal Hegde', 'Pushpa Kushal', 'Pushpa Hegde',
    'Rajesh Kushal', 'Rajesh Hegde', 'Rohit Kushal', 'Rohit Hegde',
    'Rakhi Girija', 'Rakhi Shetty', 'Girija Shetty',
    'Sarthak Malvadkar', 'Sandesh Bhagwat', 'Amod Joshi', 'Ganesh Prasad',
    'Karunakar', 'Narayna B. Shetty', 'Jayaram N. Shetty',
    'kshinternational', 'KSH International',
    'cs.connect@', 'ksh.ipo@', 'ksh@icici', 'hingnetare@', 'parag.pansare@',
    'prakash.boricha@', 'sheetal.parab@',
    'siddharth.jadhav@', 'sachin.gawade@', 'eric.bacha@',
    'tushar.gavankar@', 'pravin.teli', 'cherag.gyara@',
    'manisha.shukla@', 'hitesh.ramani@', 'sharmila.joshi@',
    'ashishmp@', 'anand.soni@bajaj',
    '45053237', '40094400', '6807 7100', '81081 14949',
    'Dhaulagiri Family', 'Everest Family Trust', 'Makalu Family',
    'Kanchenjunga Family', 'Annapurna Family', 'Broad Family Trust',
    'Waterloo Industrial', 'Kirtane & Pandit',
    'Village Birdewadi', 'Bhandary Metal',
    'U28129PN1979PLC141032',
    'Nuvama Wealth Management', 'ICICI Securities Limited',
    'Kushal Motors',
]

# Also check partial surname leaks with context
surname_checks = ['Hegde', 'Shetty', 'Bhandary', 'Malvadkar', 'Bhagwat']

full_text = ''
for p in doc.paragraphs:
    full_text += p.text + '\n'
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            full_text += c.text + '\n'

print('=== PII LEAK CHECK (Exact Matches) ===')
found_any = False
for orig in originals:
    count = full_text.lower().count(orig.lower())
    if count > 0:
        found_any = True
        idx = full_text.lower().find(orig.lower())
        context = full_text[max(0,idx-40):idx+len(orig)+40].replace('\n', ' ')
        print(f'  LEAK: "{orig}" x{count}')
        print(f'    Context: ...{context}...')

if not found_any:
    print('  All exact PII matches have been redacted!')

print('\n=== SURNAME LEAK CHECK ===')
for surname in surname_checks:
    count = full_text.count(surname)
    if count > 0:
        # Find all occurrences with context
        idx = 0
        instances = []
        while True:
            idx = full_text.find(surname, idx)
            if idx == -1:
                break
            context = full_text[max(0,idx-30):idx+len(surname)+30].replace('\n', ' ')
            instances.append(context)
            idx += len(surname)
        print(f'  "{surname}" appears {count} times:')
        for inst in instances[:3]:
            print(f'    ...{inst}...')

print(f'\nTotal document text length: {len(full_text)} chars')
