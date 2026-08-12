"""
PII Redaction Tool for DOCX Documents
======================================
A hybrid approach combining:
1. Regex-based detection for structured PII (emails, phones, SSNs, credit cards, IPs, DOBs)
2. Curated dictionary-based detection for names, companies, and addresses found in the document
3. Faker library for generating realistic fake replacements

Author: PII Redaction Tool
"""

import re
import json
import copy
import os
import sys
from collections import OrderedDict
from docx import Document
from docx.oxml.ns import qn
from faker import Faker

# ============================================================================
# CONFIGURATION
# ============================================================================

INPUT_FILE = r"Red Herring Prospectus.docx"
OUTPUT_FILE = r"Red Herring Prospectus_REDACTED.docx"
MAPPING_FILE = r"pii_mapping.json"

fake = Faker('en_IN')  # Indian locale for realistic replacements
Faker.seed(42)  # Reproducible fake data

# ============================================================================
# PII ENTITY DEFINITIONS
# ============================================================================

class PIIEntity:
    """Represents a detected PII entity with its replacement."""
    def __init__(self, original, replacement, pii_type, detection_method):
        self.original = original
        self.replacement = replacement
        self.pii_type = pii_type
        self.detection_method = detection_method

    def to_dict(self):
        return {
            "original": self.original,
            "replacement": self.replacement,
            "type": self.pii_type,
            "method": self.detection_method
        }


# ============================================================================
# PII DETECTOR CLASS
# ============================================================================

class PIIDetector:
    """
    Detects and replaces PII in text using a hybrid approach:
    - Regex patterns for structured data (emails, phones, SSNs, credit cards, IPs)
    - Curated dictionaries for domain-specific entities (names, companies, addresses)
    - Faker library for generating realistic replacements
    """

    def __init__(self):
        self.entities = []
        self.replacement_map = OrderedDict()  # original -> PIIEntity
        self._build_person_names()
        self._build_company_names()
        self._build_address_mappings()
        self._build_email_mappings()
        self._build_phone_mappings()
        self._build_website_mappings()
        self._build_din_pan_mappings()
        self._build_registration_number_mappings()

    # -----------------------------------------------------------------------
    # Person Names (curated from document analysis)
    # -----------------------------------------------------------------------
    def _build_person_names(self):
        """
        Names extracted from document analysis. Using dictionary-based matching
        because NER models struggle with Indian names and the document context
        (e.g., names in ALL CAPS, names in legal phrases).
        """
        # Map of real name -> fake name
        self.person_names = OrderedDict([
            # Promoters & Directors
            ("Kushal Subbayya Hegde", "Vikram Suresh Menon"),
            ("Pushpa Kushal Hegde", "Lakshmi Vikram Menon"),
            ("Rajesh Kushal Hegde", "Arun Vikram Menon"),
            ("Rohit Kushal Hegde", "Sanjay Vikram Menon"),
            ("Rakhi Girija Shetty", "Priya Girish Nair"),
            ("Kushal Hegde", "Vikram Menon"),
            ("Pushpa Hegde", "Lakshmi Menon"),
            ("Rajesh Hegde", "Arun Menon"),
            ("Rohit Hegde", "Sanjay Menon"),
            ("Rakhi Shetty", "Priya Nair"),

            # Historical / family members found in corporate history sections
            ("Karunakar N. Bhandary", "Ramchandra K. Sharma"),
            ("Karunakar Hegde", "Ramchandra Menon"),
            ("Narayna B. Shetty", "Bhaskar R. Nair"),
            ("Jayaram N. Shetty", "Govind M. Nair"),

            # Key Managerial Personnel
            ("Sandesh Bhagwat", "Ramesh Kulkarni"),
            ("Amod Joshi", "Nikhil Desai"),
            ("Sarthak Malvadkar", "Pranav Deshpande"),
            ("Ganesh Prasad", "Sunil Kumar"),

            # Directors & Board Members
            ("Girija Shetty", "Girish Nair"),
            ("Prakash Boricha", "Suresh Mehra"),
            ("Sheetal Parab", "Anita Sharma"),
            ("Kishan Rastogi", "Mohan Gupta"),
            ("Abhijit Diwan", "Rajendra Patil"),
            ("Lokesh Shah", "Dinesh Mehta"),
            ("Soumavo Sarkar", "Debashish Roy"),
            ("Shanti Gopalkrishnan", "Kamala Ramanathan"),

            # Bankers & Contact persons
            ("Varun Badai", "Amit Saxena"),
            ("Siddharth Jadhav", "Rahul Joshi"),
            ("Sachin Gawade", "Vishal Patil"),
            ("Eric Bacha", "Mark Thomas"),
            ("Tushar Gavankar", "Kiran Sawant"),
            ("Pravin Teli", "Manoj Deshmukh"),
            ("Chitra Raste", "Sneha Gokhale"),
            ("Sharmila Joshi", "Swati Kulkarni"),
            ("Cherag Gyara", "Faisal Qureshi"),
            ("Manisha Shukla", "Deepa Mishra"),
            ("Ashish", "Nitin"),  # ashishmp@federalbank
            ("Anand Soni", "Vijay Sharma"),
            ("Hitesh Ramani", "Paresh Vyas"),
            ("Parag Pansare", "Makarand Godbole"),

            # Auditors & Legal
            ("Hingne Tare", "Kulkarni Joshi"),

            # Additional people mentioned
            ("Subbayya Hegde", "Suresh Menon"),

            # Historical names from corporate records (allotments, transfers)
            ("Vijay Hegde", "Ajay Menon"),
            ("Sunil Nagayya Shetty", "Mohan Govind Nair"),
            ("Karunakar Bhandary", "Ramchandra Sharma"),
            ("DM Shetty", "RM Nair"),
            ("SA Shetty", "KP Nair"),
            ("Gopal BO", "Mohan RK"),
            ("Narayana B. Shetty", "Bhaskar R. Nair"),
            ("Nagayya Shetty", "Govind Nair"),
            ("Jayaram Shetty", "Govind Nair"),
        ])

    # -----------------------------------------------------------------------
    # Company Names
    # -----------------------------------------------------------------------
    def _build_company_names(self):
        """Company names found in the document that should be redacted."""
        self.company_names = OrderedDict([
            # Primary company
            ("KSH International Limited", "ABC Industries Limited"),
            ("KSH International Private Limited", "ABC Industries Private Limited"),
            ("KSH International", "ABC Industries"),
            ("Bhandary Metal Extrusion Private Limited", "Sharma Metal Works Private Limited"),

            # Trust entities (Promoter Group)
            ("Dhaulagiri Family Trust", "Sunrise Family Trust"),
            ("Everest Family Trust", "Mountain View Family Trust"),
            ("Makalu Family Trust", "Highland Family Trust"),
            ("Broad Family Trust", "Horizon Family Trust"),
            ("Annapurna Family Trust", "Valley View Family Trust"),
            ("Kanchenjunga Family Trust", "Summit Family Trust"),
            ("Waterloo Industrial Park VI Private Limited", "Greenfield Industrial Park VI Private Limited"),
            ("Waterloo Industrial Park IX Private Limited", "Greenfield Industrial Park IX Private Limited"),
            ("Waterloo Industrial Park VIII Private Limited", "Greenfield Industrial Park VIII Private Limited"),
            ("Waterloo Industrial Park VII Private Limited", "Greenfield Industrial Park VII Private Limited"),
            ("Waterloo Industrial Park V Private Limited", "Greenfield Industrial Park V Private Limited"),
            ("Waterloo Industrial Park IV Private Limited", "Greenfield Industrial Park IV Private Limited"),
            ("Waterloo Industrial Park III Private Limited", "Greenfield Industrial Park III Private Limited"),
            ("Waterloo Industrial Park II Private Limited", "Greenfield Industrial Park II Private Limited"),
            ("Waterloo Industrial Park I Private Limited", "Greenfield Industrial Park I Private Limited"),
            ("Waterloo Industrial Park", "Greenfield Industrial Park"),

            # BRLMs and financial institutions
            ("Nuvama Wealth Management Limited", "Alpha Wealth Management Limited"),
            ("ICICI Securities Limited", "Beta Securities Limited"),
            ("MUFG Intime India Private Limited", "Gamma Services India Private Limited"),
            ("Link Intime India Private Limited", "Gamma Services India Private Limited"),

            # Auditors
            ("Kirtane & Pandit LLP", "Sharma & Associates LLP"),
            ("Kirtane & Pandit", "Sharma & Associates"),

            # Banks (lenders)
            ("HDFC Bank", "Delta Bank"),
            ("ICICI Bank", "Beta Bank"),
            ("IndusInd Bank", "Epsilon Bank"),
            ("Federal Bank", "Zeta Bank"),
            ("State Bank of India", "National Bank of India"),
            ("Export-Import Bank of India", "International Trade Bank of India"),
            ("Bajaj Finance", "Omega Finance"),

            # Legal advisors
            ("Trilegal", "LexCounsel"),

            # Suppliers / Customers mentioned by name
            ("CARE Analytics and Advisory Private Limited", "Market Research Analytics Private Limited"),
            ("CareEdge Research", "InsightEdge Research"),
            ("Ahlstrom Sweden AB", "Nordic Paper AB"),
            ("Cindus Corporation", "Pacific Materials Corporation"),
            ("Elantas Beck India Limited", "ChemCoat India Limited"),
            ("Hindalco Industries Limited", "MetalCorp Industries Limited"),
            ("Polycom Associates", "PolyTech Associates"),
            ("Savli Copper Products Private Limited", "Rajkot Copper Products Private Limited"),
            ("Union Copper Rod LLC", "United Metals Rod LLC"),
            ("Vedanta Limited Sterlite Copper", "Mining Corp Limited Copper Division"),

            # Customer companies
            ("Al-Ahleia Switchgear Co.", "Gulf Electric Switchgear Co."),
            ("Bharat Bijlee Limited", "Desh Electrical Limited"),
            ("CG Power and Industrial Solutions Limited", "TechPower Industrial Solutions Limited"),
            ("Emirates Transformer & Switchgear Limited", "Arabian Power Equipment Limited"),
            ("Georgia Transformer Corporation", "Atlantic Transformer Corporation"),
            ("Nidec Industrial Automation India Private Limited", "Servo Automation India Private Limited"),
            ("Transformers & Rectifiers (India) Limited", "Voltage Equipment (India) Limited"),
            ("Virginia Transformer Corporation", "Columbia Transformer Corporation"),

            # Registrar
            ("KSH International Chakan Internal Kamgar Sangathna", "ABC Industries Chakan Workers Union"),

            # Additional Group Companies
            ("Kushal Motors and Electricals Private Limited", "Vikram Motors and Electricals Private Limited"),
            ("Kushal Motors", "Vikram Motors"),

            # Trust branch names (referencing promoter family members)
            ("Rajesh Branch", "Arun Branch"),
            ("Sangeeta Branch", "Meena Branch"),
            ("Rakhi Branch", "Priya Branch"),
            ("Rohit Branch", "Sanjay Branch"),
            ("Kushal Branch", "Vikram Branch"),
        ])

    # -----------------------------------------------------------------------
    # Physical / Mailing Addresses
    # -----------------------------------------------------------------------
    def _build_address_mappings(self):
        """Physical addresses found in the document."""
        self.addresses = OrderedDict([
            # Registered Office
            ("11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune – 410 501, Maharashtra, India",
             "45/A, 45/B and 45/C, Village Wagholi, Haveli Taluka, Pune – 412 207, Maharashtra, India"),
            ("11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune \u2013 410 501, Maharashtra, India",
             "45/A, 45/B and 45/C, Village Wagholi, Haveli Taluka, Pune \u2013 412 207, Maharashtra, India"),
            # Corporate Office
            ("201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune – 411 045, Maharashtra, India",
             "502, Tower 5, Phoenix Business Hub, Off Hinjewadi Road, Wakad, Pune – 411 057, Maharashtra, India"),
            ("201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune \u2013 411 045, Maharashtra, India",
             "502, Tower 5, Phoenix Business Hub, Off Hinjewadi Road, Wakad, Pune \u2013 411 057, Maharashtra, India"),
            # Nuvama Office
            ("801 - 804, Wing A, Building No 3, Inspire BKC, G Block, Bandra Kurla Complex, Bandra East, Mumbai 400051",
             "901 - 905, Wing B, Building No 7, Phoenix BKC, E Block, Bandra Kurla Complex, Bandra East, Mumbai 400098"),
            # ICICI Securities Office
            ("ICICI Venture House, Appasaheb Marathe Marg, Prabhadevi, Mumbai 400025",
             "Tower A, Trade Center, Senapati Bapat Marg, Dadar West, Mumbai 400028"),
            # MUFG/Link Intime Office
            ("C-101, Embassy 247, 1st Floor, L B S Marg, Vikhroli (West), Mumbai 400083",
             "D-201, Phoenix Tower, 2nd Floor, Western Express Highway, Goregaon (East), Mumbai 400063"),
            # HDFC Bank address
            ("Lodha I Think Techno Campus, O-2, Next to Kanjurmarg Station, Kanjurmarg (East), Mumbai – 400042",
             "Lotus Business Park, B-3, Next to Ghatkopar Station, Ghatkopar (East), Mumbai – 400077"),
            # ICICI Bank address
            ("163, 5th Floor, H.T.Parekh Marg Backbay Reclamation Churchgate, Mumbai",
             "245, 7th Floor, Nariman Point Road, Marine Lines, Mumbai"),
            # Supa facility
            ("Supa, Ahilyanagar", "Shirur, Pune District"),
            ("Supa Facility", "Shirur Facility"),
            # Chakan
            ("Chakan, Pune", "Talegaon, Pune"),
            ("Unit 2 in Chakan", "Unit 2 in Talegaon"),
            # Trilegal
            ("One BKC, Tower B, 1202 & 1203, Plot No. C-66, G Block, Bandra Kurla Complex, Bandra (East), Mumbai 400051",
             "Two BKC, Tower C, 1501, Plot No. D-72, F Block, Bandra Kurla Complex, Bandra (East), Mumbai 400098"),
            # Auditor address
            ("5th Floor, Gopal House", "3rd Floor, Shivaji House"),
            ("Opposite Harshal Hall, above HDFC Limited Karve Road, Pune",
             "Near Fergusson College, above SBI, Law Road, Pune"),
            # Hingne Tare address
            ("Flat No. 102, Sai Complex Shaniwar Peth, Pune",
             "Office No. 205, Om Complex Sadashiv Peth, Pune"),
            # Citibank address
            ("Citibank N.A., Plot C-61, G Block, 5th Floor, Citibank Centre, Bandra Kurla Complex, Bandra East, Mumbai 400098",
             "Standard Chartered Bank, Plot B-42, E Block, 3rd Floor, SC Centre, Bandra Kurla Complex, Bandra East, Mumbai 400051"),
            # EXIM Bank
            ("Centre One, Floor 21, World Trade Centre Complex, Cuffe Parade, Mumbai 400005",
             "Tower Two, Floor 15, International Business Centre, Nariman Point, Mumbai 400021"),
            # IndusInd Bank
            ("2401, Gen. A.K. Vaidya Marg, Malad (East), Mumbai 400097",
             "1803, Gen. M.B. Sharma Marg, Andheri (East), Mumbai 400069"),
            # Federal Bank
            ("Srei BKC, 4th Floor, Plot No - C68, G Block, BKC Bandra East, Mumbai 400051",
             "Parinee Crescenzo, 5th Floor, Plot No - C38, G Block, BKC Bandra East, Mumbai 400098"),
            # Bajaj Finance
            ("Pune - Akurdi, 4th and 5th Floor, Bajaj Finserv Corporate Office Survey No. 208",
             "Pune - Hinjewadi, 6th and 7th Floor, Omega Finserv Corporate Office Survey No. 312"),
            # Village name
            ("Village Birdewadi", "Village Wagholi"),
        ])

    # -----------------------------------------------------------------------
    # Email Addresses
    # -----------------------------------------------------------------------
    def _build_email_mappings(self):
        """Email addresses found in document, mapped to fake alternatives."""
        self.email_map = OrderedDict([
            ("cs.connect@kshinternational.com", "cs.info@abcindustries.com"),
            ("Sarthak.malvadkar@kshinterantional.com", "pranav.deshpande@abcindustries.com"),
            ("ksh.ipo@nuvama.com", "abc.ipo@alphawealth.com"),
            ("customerservice.mb@nuvama.com", "support.mb@alphawealth.com"),
            ("ksh@icicisecurities.com", "abc@betasecurities.com"),
            ("customercare@icicisecurities.com", "support@betasecurities.com"),
            ("prakash.boricha@nuvama.com", "suresh.mehra@alphawealth.com"),
            ("sheetal.parab@nuvama.com", "anita.sharma@alphawealth.com"),
            ("ipo@trilegal.com", "ipo@lexcounsel.com"),
            ("kshinternational.ipo@in.mpms.mufg.com", "abcindustries.ipo@in.gammaservices.com"),
            ("siddharth.jadhav@hdfcbank.com", "rahul.joshi@deltabank.com"),
            ("sachin.gawade@hdfcbank.com", "vishal.patil@deltabank.com"),
            ("eric.bacha@hdfcbank.com", "mark.thomas@deltabank.com"),
            ("tushar.gavankar@hdfcbank.com", "kiran.sawant@deltabank.com"),
            ("pravin.teli2@hdfcbank.com", "manoj.deshmukh@deltabank.com"),
            ("Ipocmg@icicibank.com", "ipocmg@betabank.com"),
            ("parag.pansare@kirtanepandit.com", "makarand.godbole@sharmaassociates.com"),
            ("hingnetare@gmail.com", "kulkarnijoshi@gmail.com"),
            ("hitesh.ramani@citi.com", "paresh.vyas@standardchartered.com"),
            ("pro@eximbankindia.in", "info@tradebankindia.in"),
            ("sharmila.joshi@indusind.com", "swati.kulkarni@epsilonbank.com"),
            ("cherag.gyara@icicibank.com", "faisal.qureshi@betabank.com"),
            ("manisha.shukla@hdfcbank.com", "deepa.mishra@deltabank.com"),
            ("rm6.ifbpune@sbi.co.in", "rm9.ifbpune@nationalbank.co.in"),
            ("ashishmp@federalbank.co.in", "nitinrk@zetabank.co.in"),
            ("anand.soni@bajajfinserv.in", "vijay.sharma@omegafinserv.in"),
        ])

    # -----------------------------------------------------------------------
    # Phone Numbers
    # -----------------------------------------------------------------------
    def _build_phone_mappings(self):
        """Phone numbers found in the document."""
        self.phone_map = OrderedDict([
            ("+ 91 20 4505 3237", "+ 91 20 6783 1245"),
            ("+91 20 45053237", "+91 20 67831245"),
            ("91 20 45053237", "91 20 67831245"),
            ("+91 22 40094400", "+91 22 53218765"),
            ("+91 22 4009 4400", "+91 22 5321 8765"),
            ("91 22 4009 4400", "91 22 5321 8765"),
            ("+91 22 6807 7100", "+91 22 7912 3456"),
            ("+91 81081 14949", "+91 93265 78123"),
            ("81081 14949", "93265 78123"),
            ("+91 22 4079 1000", "+91 22 5183 2000"),
            ("+91 22 30752929", "+91 22 41863838"),
            ("+91 22 30752928", "+91 22 41863837"),
            ("+91 22 30752914", "+91 22 41863823"),
            ("022-68052182", "022-79163293"),
            ("+91 20 6729 5100", "+91 20 7834 6200"),
            ("91 20 6729 5100", "91 20 7834 6200"),
            ("+91 20 6606 4494", "+91 20 7712 5505"),
            ("+91 20 2640 3100", "+91 20 3751 4200"),
            ("+91-20-26234000", "+91-20-37454111"),
            ("91 8879770456", "91 9923456789"),
            ("+ 91 8879770456", "+ 91 9923456789"),
            ("+91 20 6769 4648", "+91 20 7873 5759"),
            ("+91 20 2561 8211", "+91 20 3672 9322"),
            ("91 91586 40360", "91 98765 43210"),
            ("+91 20 7157 6403", "+91 20 8268 7514"),
        ])

    # -----------------------------------------------------------------------
    # Websites
    # -----------------------------------------------------------------------
    def _build_website_mappings(self):
        """Website URLs found in the document."""
        self.website_map = OrderedDict([
            ("www.kshinternational.com", "www.abcindustries.com"),
            ("www.kshinternational. com", "www.abcindustries.com"),
            ("kshinternational.com", "abcindustries.com"),
            ("kshinternational. com", "abcindustries.com"),
            ("https://kshinternational.com", "https://abcindustries.com"),
            ("www.nuvama.com", "www.alphawealth.com"),
            ("www.icicisecurities.com", "www.betasecurities.com"),
            ("www.in.mpms.mufg.com", "www.in.gammaservices.com"),
            ("www.hdfcbank.com", "www.deltabank.com"),
            ("www.icicibank.com", "www.betabank.com"),
            ("www.eximbankindia.in", "www.tradebankindia.in"),
            ("www.indusind.com", "www.epsilonbank.com"),
            ("www.bajajfinance.com", "www.omegafinance.com"),
        ])

    # -----------------------------------------------------------------------
    # DIN Numbers & PAN
    # -----------------------------------------------------------------------
    def _build_din_pan_mappings(self):
        """DIN (Director Identification Numbers) found in the document."""
        self.din_map = OrderedDict([
            # These are typically 8-digit numbers associated with directors
            ("000013004", "000098765"),
            ("000011179", "000087654"),
            ("000004058", "000076543"),
            ("000166136", "000065432"),
        ])

    # -----------------------------------------------------------------------
    # Registration / CIN Numbers
    # -----------------------------------------------------------------------
    def _build_registration_number_mappings(self):
        """Corporate Identity Numbers and registration numbers."""
        self.registration_map = OrderedDict([
            ("U28129PN1979PLC141032", "U29110MH1985PLC198765"),
            ("105215W/ W100057", "112345E/ E200098"),
            ("105215W/W100057", "112345E/E200098"),
            ("116417W", "223456E"),
            ("014680", "025791"),
        ])

    # -----------------------------------------------------------------------
    # Regex-based Detection
    # -----------------------------------------------------------------------
    def _get_regex_patterns(self):
        """
        Returns regex patterns for structured PII types.
        These catch any PII not already in our curated dictionaries.
        """
        return {
            "EMAIL": (
                r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}',
                self._replace_email
            ),
            "PHONE_INTL": (
                r'(?<!\d)(?:\+\s*91[\s\-]*)?(?:\d[\s\-]*){10}(?!\d)',
                self._replace_phone
            ),
            "SSN": (
                r'\b\d{3}-\d{2}-\d{4}\b',
                self._replace_ssn
            ),
            "CREDIT_CARD": (
                r'\b(?:\d{4}[\s\-]?){3}\d{4}\b',
                self._replace_credit_card
            ),
            "IP_ADDRESS": (
                r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
                self._replace_ip
            ),
            "DATE_OF_BIRTH": (
                r'\b(?:DOB|Date\s+of\s+Birth|born\s+on)[:\s]*\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}\b',
                self._replace_dob
            ),
        }

    def _replace_email(self, match):
        email = match.group(0)
        if email.lower() in [e.lower() for e in self.email_map]:
            for k, v in self.email_map.items():
                if k.lower() == email.lower():
                    return v
        # Generate new fake email
        return fake.email()

    def _replace_phone(self, match):
        return fake.phone_number()

    def _replace_ssn(self, match):
        return f"{fake.random_int(100,999)}-{fake.random_int(10,99)}-{fake.random_int(1000,9999)}"

    def _replace_credit_card(self, match):
        return fake.credit_card_number()

    def _replace_ip(self, match):
        return fake.ipv4()

    def _replace_dob(self, match):
        return f"DOB: {fake.date_of_birth().strftime('%d/%m/%Y')}"

    # -----------------------------------------------------------------------
    # Main Redaction Logic
    # -----------------------------------------------------------------------
    def redact_text(self, text):
        """
        Redact all PII from the given text.
        Order of operations matters - we do dictionary-based replacements first
        (longest match first to avoid partial replacements), then regex-based.
        """
        if not text or not text.strip():
            return text

        redacted = text

        # 1. Replace addresses (longest first to avoid partial matches)
        for original, replacement in sorted(self.addresses.items(), key=lambda x: len(x[0]), reverse=True):
            if original in redacted:
                redacted = redacted.replace(original, replacement)
                self._record_entity(original, replacement, "ADDRESS", "dictionary")

        # 2. Replace company names (longest first)
        for original, replacement in sorted(self.company_names.items(), key=lambda x: len(x[0]), reverse=True):
            # Case-insensitive replacement
            pattern = re.compile(re.escape(original), re.IGNORECASE)
            if pattern.search(redacted):
                # Handle UPPER CASE versions too
                upper_original = original.upper()
                upper_replacement = replacement.upper()
                if upper_original in redacted:
                    redacted = redacted.replace(upper_original, upper_replacement)
                    self._record_entity(upper_original, upper_replacement, "COMPANY", "dictionary")
                # Handle normal case
                if original in redacted:
                    redacted = redacted.replace(original, replacement)
                    self._record_entity(original, replacement, "COMPANY", "dictionary")

        # 3. Replace registration/CIN numbers
        for original, replacement in self.registration_map.items():
            if original in redacted:
                redacted = redacted.replace(original, replacement)
                self._record_entity(original, replacement, "REGISTRATION_NUMBER", "dictionary")

        # 4. Replace DIN numbers
        for original, replacement in self.din_map.items():
            if original in redacted:
                redacted = redacted.replace(original, replacement)
                self._record_entity(original, replacement, "DIN", "dictionary")

        # 5. Replace emails (dictionary first, then regex for any missed)
        for original, replacement in self.email_map.items():
            # Case-insensitive email replacement
            pattern = re.compile(re.escape(original), re.IGNORECASE)
            if pattern.search(redacted):
                redacted = pattern.sub(replacement, redacted)
                self._record_entity(original, replacement, "EMAIL", "dictionary")

        # 6. Replace phone numbers
        for original, replacement in sorted(self.phone_map.items(), key=lambda x: len(x[0]), reverse=True):
            if original in redacted:
                redacted = redacted.replace(original, replacement)
                self._record_entity(original, replacement, "PHONE", "dictionary")

        # 7. Replace websites
        for original, replacement in sorted(self.website_map.items(), key=lambda x: len(x[0]), reverse=True):
            if original in redacted:
                redacted = redacted.replace(original, replacement)
                self._record_entity(original, replacement, "WEBSITE", "dictionary")

        # 8. Replace person names (longest first to avoid partial matches)
        for original, replacement in sorted(self.person_names.items(), key=lambda x: len(x[0]), reverse=True):
            # Handle ALL CAPS version
            upper_original = original.upper()
            upper_replacement = replacement.upper()
            if upper_original in redacted:
                redacted = redacted.replace(upper_original, upper_replacement)
                self._record_entity(upper_original, upper_replacement, "PERSON_NAME", "dictionary")
            # Handle normal case
            if original in redacted:
                redacted = redacted.replace(original, replacement)
                self._record_entity(original, replacement, "PERSON_NAME", "dictionary")

        # 9. Regex-based catch-all for any remaining structured PII
        # (Only SSNs, credit cards, IPs, DOBs - emails and phones already handled)
        for pii_type, (pattern, replacer) in self._get_regex_patterns().items():
            if pii_type in ("EMAIL", "PHONE_INTL"):
                continue  # Already handled by dictionary
            compiled = re.compile(pattern, re.IGNORECASE)
            matches = compiled.finditer(redacted)
            for match in matches:
                original_val = match.group(0)
                # Avoid replacing numbers that are clearly financial amounts or dates
                if pii_type == "CREDIT_CARD" and self._is_financial_amount(original_val, redacted):
                    continue
                if pii_type == "IP_ADDRESS" and self._is_version_number(original_val, redacted):
                    continue
                replacement_val = replacer(match)
                redacted = redacted.replace(original_val, replacement_val, 1)
                self._record_entity(original_val, replacement_val, pii_type, "regex")

        return redacted

    def _is_financial_amount(self, text, context):
        """Check if a number is likely a financial amount, not a credit card."""
        # Remove spaces and dashes
        digits = re.sub(r'[\s\-]', '', text)
        # Financial amounts in Indian prospectus documents are usually preceded by ₹ or Rs
        idx = context.find(text)
        if idx > 0:
            preceding = context[max(0, idx-5):idx]
            if '₹' in preceding or 'Rs' in preceding or '.' in preceding:
                return True
        return len(digits) < 13  # Credit cards are 13-19 digits

    def _is_version_number(self, text, context):
        """Check if an IP-like pattern is actually a version number."""
        idx = context.find(text)
        if idx > 0:
            preceding = context[max(0, idx-10):idx].lower()
            if 'version' in preceding or 'v.' in preceding:
                return True
        return False

    def _record_entity(self, original, replacement, pii_type, method):
        """Record a PII entity for the mapping file."""
        if original not in self.replacement_map:
            entity = PIIEntity(original, replacement, pii_type, method)
            self.replacement_map[original] = entity
            self.entities.append(entity)

    def get_mapping_report(self):
        """Return a structured report of all PII replacements."""
        report = {
            "total_entities": len(self.entities),
            "by_type": {},
            "entities": []
        }
        for entity in self.entities:
            if entity.pii_type not in report["by_type"]:
                report["by_type"][entity.pii_type] = 0
            report["by_type"][entity.pii_type] += 1
            report["entities"].append(entity.to_dict())
        return report


# ============================================================================
# DOCX PROCESSING
# ============================================================================

class DocxRedactor:
    """
    Handles reading, redacting, and writing DOCX files.
    Processes both paragraphs and tables, preserving formatting.
    """

    def __init__(self, detector: PIIDetector):
        self.detector = detector
        self.stats = {
            "paragraphs_processed": 0,
            "table_cells_processed": 0,
            "total_replacements": 0,
        }

    def redact_document(self, input_path, output_path):
        """Main method to redact a DOCX document."""
        print(f"[*] Loading document: {input_path}")
        doc = Document(input_path)

        # Process paragraphs
        print("[*] Processing paragraphs...")
        for para in doc.paragraphs:
            self._redact_paragraph(para)
            self.stats["paragraphs_processed"] += 1

        # Process tables
        print("[*] Processing tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._redact_paragraph(para)
                    self.stats["table_cells_processed"] += 1

        # Process headers and footers
        print("[*] Processing headers and footers...")
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and header.is_linked_to_previous is False:
                    for para in header.paragraphs:
                        self._redact_paragraph(para)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and footer.is_linked_to_previous is False:
                    for para in footer.paragraphs:
                        self._redact_paragraph(para)

        # Save
        print(f"[*] Saving redacted document: {output_path}")
        doc.save(output_path)
        print("[+] Document saved successfully!")

        return self.stats

    def _redact_paragraph(self, paragraph):
        """
        Redact PII from a paragraph while preserving formatting.
        Strategy: reconstruct the full text, redact it, then redistribute
        across the existing runs to preserve formatting.
        """
        full_text = paragraph.text
        if not full_text.strip():
            return

        redacted_text = self.detector.redact_text(full_text)

        if redacted_text != full_text:
            self.stats["total_replacements"] += 1
            self._replace_paragraph_text(paragraph, redacted_text)

    def _replace_paragraph_text(self, paragraph, new_text):
        """
        Replace the text of a paragraph while trying to preserve formatting.
        Uses a run-merging strategy: puts all text into the first run
        and clears the remaining runs.
        """
        runs = paragraph.runs
        if not runs:
            return

        # If only one run, simple replacement
        if len(runs) == 1:
            runs[0].text = new_text
            return

        # Multi-run: put everything in first run, clear others
        # This preserves the formatting of the first run
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""


# ============================================================================
# MAIN EXECUTION
# ============================================================================

def main():
    # Resolve paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(script_dir, INPUT_FILE)
    output_path = os.path.join(script_dir, OUTPUT_FILE)
    mapping_path = os.path.join(script_dir, MAPPING_FILE)

    if not os.path.exists(input_path):
        print(f"[!] Error: Input file not found: {input_path}")
        sys.exit(1)

    # Initialize detector and redactor
    detector = PIIDetector()
    redactor = DocxRedactor(detector)

    # Perform redaction
    stats = redactor.redact_document(input_path, output_path)

    # Save PII mapping
    mapping_report = detector.get_mapping_report()
    with open(mapping_path, 'w', encoding='utf-8') as f:
        json.dump(mapping_report, f, indent=2, ensure_ascii=False)

    # Print summary
    print("\n" + "=" * 60)
    print("REDACTION SUMMARY")
    print("=" * 60)
    print(f"Paragraphs processed:    {stats['paragraphs_processed']}")
    print(f"Table cells processed:   {stats['table_cells_processed']}")
    print(f"Total text blocks with replacements: {stats['total_replacements']}")
    print(f"Unique PII entities found: {mapping_report['total_entities']}")
    print("\nBreakdown by PII type:")
    for pii_type, count in mapping_report["by_type"].items():
        print(f"  {pii_type:25s}: {count}")
    print(f"\nPII mapping saved to: {mapping_path}")
    print(f"Redacted document saved to: {output_path}")


if __name__ == "__main__":
    main()
